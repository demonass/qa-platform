from fastapi import FastAPI, HTTPException, File, UploadFile, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sse_starlette.sse import EventSourceResponse
import asyncio
import json
import os
from typing import Optional
from agent_service import conversation_history, run_agent_with_history, run_agent_with_tools, rag_service, stream_chat_response
from document_processor import process_document, DocumentProcessor
from config import LLM_PROVIDER, get_llm_config, get_redis_config

# Redis 缓存服务
try:
    from redis_cache import cache_service
except ImportError:
    cache_service = None

app = FastAPI(title="QA Agent Service", version="1.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup_event():
    """应用启动时初始化 RAG 服务"""
    print("[INFO] Initializing RAG service on startup...")
    rag_service.initialize()


class ChatRequest(BaseModel):
    message: str
    session_id: str = "default"
    mode: str = "default"

class ChatResponse(BaseModel):
    response: str
    status: str

class ExportRequest(BaseModel):
    session_id: str = "default"
    format: str = "markdown"

class IntentRequest(BaseModel):
    message: str

class LoginRequest(BaseModel):
    username: str
    password: str

class RegisterRequest(BaseModel):
    username: str
    password: str
    role: str = "user"

class ChangePasswordRequest(BaseModel):
    user_id: str
    new_password: str

class SessionListRequest(BaseModel):
    user_id: Optional[str] = None
    is_admin: bool = False


def _get_user_id(x_user_id: Optional[str] = Header(None)) -> Optional[str]:
    """Extract user_id from X-User-Id header (set by Go backend after JWT verification)"""
    return x_user_id

def _get_user_role(x_user_role: Optional[str] = Header(None)) -> Optional[str]:
    """Extract role from X-User-Role header"""
    return x_user_role

def _is_admin(x_user_role: Optional[str] = Header(None)) -> bool:
    return x_user_role == "admin"


@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "qa-agent"}


@app.get("/config")
async def get_config():
    llm_config = get_llm_config()
    redis_config = get_redis_config()
    return {
        "provider": LLM_PROVIDER,
        "model": llm_config["model"],
        "api_base": llm_config["api_base"],
        "version": "2.0.0",
        "cache": {
            "enabled": redis_config["enabled"],
            "ttl_seconds": redis_config["ttl_seconds"],
            "host": redis_config["host"],
            "port": redis_config["port"]
        }
    }


@app.post("/intent")
async def detect_intent(request: IntentRequest):
    from agent_service import Intent, detect_intent
    from intent_detector import DETECTION_METHOD, DETECTION_METHOD_DESCRIPTION, intent_classifier

    try:
        intent = detect_intent(request.message)
        intent_name = intent.name if isinstance(intent, Intent) else str(intent)

        intent_descriptions = {
            "TEST_CASE": "测试用例生成",
            "TEST_PLAN": "测试计划制定",
            "CODE_ANALYSIS": "代码分析",
            "RAG_QA": "知识库问答",
            "RUN_TESTS": "测试执行",
            "DEFAULT": "默认回答"
        }

        description = intent_descriptions.get(intent_name, "未知意图")

        from intent_detector import INTENT_SAMPLES

        result = {
            "intent": intent_name,
            "description": description,
            "detection_method": DETECTION_METHOD,
            "detection_method_description": DETECTION_METHOD_DESCRIPTION,
            "message": f"输入已识别为「{description}」意图"
        }

        if intent_classifier is not None:
            text_embedding = intent_classifier.embed_text(request.message)
            max_sim = 0
            target_intent = Intent[intent_name] if intent_name in Intent.__members__ else None
            if target_intent and target_intent in INTENT_SAMPLES:
                for sample in INTENT_SAMPLES[target_intent]:
                    sample_embedding = intent_classifier.embed_text(sample)
                    sim = intent_classifier.cosine_similarity(text_embedding, sample_embedding)
                    if sim > max_sim:
                        max_sim = sim
            result["confidence"] = round(max_sim, 4)
            result["model_used"] = "bge-base-zh-v1.5"

        return result
    except Exception as e:
        return {
            "intent": "UNKNOWN",
            "description": "未知意图",
            "detection_method": DETECTION_METHOD,
            "detection_method_description": DETECTION_METHOD_DESCRIPTION,
            "message": f"意图检测失败: {str(e)}"
        }


@app.post("/chat")
async def chat(request: ChatRequest, x_user_id: Optional[str] = Header(None)):
    try:
        conversation_history.add_message(request.session_id, "user", request.message, user_id=x_user_id)

        result = run_agent_with_tools(request.session_id, request.message)

        conversation_history.add_message(request.session_id, "assistant", result, user_id=x_user_id)

        return ChatResponse(
            response=result,
            status="success"
        )
    except Exception as e:
        return ChatResponse(
            response=f"Error: {str(e)}",
            status="error"
        )


@app.post("/chat/stream")
async def chat_stream(request: ChatRequest, x_user_id: Optional[str] = Header(None)):
    """真正的 token 级流式对话端点"""
    conversation_history.add_message(request.session_id, "user", request.message, user_id=x_user_id)

    async def event_generator():
        try:
            async for token in stream_chat_response(request.session_id, request.message, request.mode, user_id=x_user_id):
                yield {
                    "event": "message",
                    "data": json.dumps({"content": token, "done": False})
                }

            yield {
                "event": "message",
                "data": json.dumps({"content": "", "done": True})
            }
        except Exception as e:
            yield {
                "event": "error",
                "data": json.dumps({"error": str(e)})
            }

    return EventSourceResponse(event_generator())


@app.get("/history/{session_id}")
async def get_history(session_id: str):
    history = conversation_history.get_history(session_id)
    return {"session_id": session_id, "history": history}


@app.delete("/history/{session_id}")
async def clear_history(session_id: str):
    conversation_history.clear_history(session_id)
    return {"status": "success", "message": f"History cleared for session: {session_id}"}


@app.post("/export")
async def export_conversation(request: ExportRequest):
    history = conversation_history.get_history(request.session_id)

    if not history:
        raise HTTPException(status_code=404, detail="No conversation history found")

    if request.format == "markdown":
        md_content = "# 对话记录导出\n\n"
        for msg in history:
            role = "👤 用户" if msg["role"] == "user" else "🤖 AI 助手"
            md_content += f"## {role}\n\n{msg['content']}\n\n---\n\n"
        return {"format": "markdown", "content": md_content}

    elif request.format == "text":
        text_content = "对话记录导出\n" + "=" * 50 + "\n\n"
        for msg in history:
            role = "用户" if msg["role"] == "user" else "AI助手"
            text_content += f"[{role}]\n{msg['content']}\n\n"
        return {"format": "text", "content": text_content}

    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'markdown' or 'text'")


class RAGQueryRequest(BaseModel):
    question: str


@app.post("/rag/query")
async def rag_query(request: RAGQueryRequest):
    if not rag_service:
        raise HTTPException(status_code=503, detail="RAG service not available")

    result = rag_service.query(request.question)
    if result["success"]:
        return {
            "answer": result["answer"],
            "sources": result["sources"],
            "status": "success"
        }
    else:
        raise HTTPException(status_code=500, detail=result["answer"])


@app.post("/rag/reload")
async def rag_reload():
    if not rag_service:
        raise HTTPException(status_code=503, detail="RAG service not available")

    try:
        rag_service.initialize()
        return {"status": "success", "message": "RAG index reloaded"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to reload RAG index: {str(e)}")


@app.get("/rag/status")
async def rag_status():
    if not rag_service:
        return {"status": "not_available", "message": "RAG service not initialized"}

    return {"status": "available", "message": "RAG service is ready"}


# ==================== Cache Endpoints ====================

@app.get("/cache/status")
async def cache_status():
    """获取缓存服务状态"""
    if not cache_service:
        return {"status": "not_available", "message": "缓存服务未加载"}
    
    return cache_service.get_stats()


@app.delete("/cache/clear")
async def cache_clear(session_id: Optional[str] = None):
    """清除缓存（可选按会话ID过滤）"""
    if not cache_service:
        raise HTTPException(status_code=503, detail="缓存服务不可用")
    
    success = cache_service.clear(session_id)
    if success:
        return {"status": "success", "message": f"缓存清除成功{'' if not session_id else f'（会话: {session_id}）'}"}
    else:
        raise HTTPException(status_code=500, detail="缓存清除失败")


# ==================== Auth Endpoints ====================

@app.post("/auth/login")
async def auth_login(request: LoginRequest):
    """Verify credentials and return user info. Called by Go backend after it issues JWT."""
    user = conversation_history.verify_user(request.username, request.password)
    if not user:
        raise HTTPException(status_code=401, detail="用户名或密码错误")
    return {"status": "success", "user": user}


@app.post("/auth/register")
async def auth_register(request: RegisterRequest, x_user_role: Optional[str] = Header(None)):
    """Create a new user. Admin can set any role; self-registration defaults to 'user'."""
    # Self-registration: force role to 'user' unless admin
    if x_user_role != "admin":
        request.role = "user"

    if not request.username or not request.password:
        raise HTTPException(status_code=400, detail="用户名和密码不能为空")
    if len(request.password) < 6:
        raise HTTPException(status_code=400, detail="密码至少6位")
    if len(request.username) < 2:
        raise HTTPException(status_code=400, detail="用户名至少2位")
    try:
        user = conversation_history.add_user(request.username, request.password, request.role)
        return {"status": "success", "user": user}
    except Exception as e:
        raise HTTPException(status_code=409, detail=f"创建用户失败: {str(e)}")


@app.get("/auth/users")
async def list_users(x_user_role: Optional[str] = Header(None)):
    """List all users (admin only)"""
    if x_user_role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可查看")
    users = conversation_history.list_users()
    return {"users": users}


@app.delete("/auth/users/{user_id}")
async def delete_user(user_id: str, x_user_role: Optional[str] = Header(None)):
    """Delete a user (admin only)"""
    if x_user_role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除")
    ok = conversation_history.delete_user(user_id)
    if not ok:
        raise HTTPException(status_code=400, detail="无法删除：不能删除最后一个管理员")
    return {"status": "success"}


@app.post("/auth/change-password")
async def change_password(request: ChangePasswordRequest):
    """Change a user's password"""
    conversation_history.change_password(request.user_id, request.new_password)
    return {"status": "success"}


@app.get("/sessions")
async def list_sessions(x_user_id: Optional[str] = Header(None), x_user_role: Optional[str] = Header(None)):
    """列出当前用户的会话（admin 可看全部）"""
    is_admin = (x_user_role == "admin")
    sessions = conversation_history.list_sessions(user_id=x_user_id, is_admin=is_admin)
    return {"sessions": sessions}


# ==================== Document Endpoints ====================

class DocumentProcessRequest(BaseModel):
    content: str
    strategy: str = "semantic"
    target_chunks: int = 5


class DocumentChunkRequest(BaseModel):
    content: str
    chunk_size: int = 500
    chunk_overlap: int = 50


@app.post("/document/process")
async def document_process(request: DocumentProcessRequest):
    try:
        if not request.content.strip():
            raise HTTPException(status_code=400, detail="文档内容不能为空")

        if request.strategy not in ["recursive", "semantic", "topic"]:
            raise HTTPException(status_code=400, detail="无效的切分策略")

        result = process_document(
            text=request.content,
            strategy=request.strategy,
            target_chunks=request.target_chunks
        )

        return {
            "status": "success",
            "strategy": request.strategy,
            "module_count": len(result),
            "modules": result
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档处理失败: {str(e)}")


@app.post("/document/chunk")
async def document_chunk(request: DocumentChunkRequest):
    try:
        if not request.content.strip():
            raise HTTPException(status_code=400, detail="文档内容不能为空")

        processor = DocumentProcessor(
            chunk_size=request.chunk_size,
            chunk_overlap=request.chunk_overlap
        )

        chunks = processor.split_by_recursive_char(request.content)

        return {
            "status": "success",
            "chunk_count": len(chunks),
            "chunks": chunks
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档切分失败: {str(e)}")


@app.post("/document/upload")
async def document_upload(
    file: UploadFile = File(...),
    strategy: str = "semantic",
    target_chunks: int = 5
):
    try:
        print(f"[INFO] Received file upload request: {file.filename}")
        
        allowed_extensions = [".txt", ".md", ".docx", ".pdf"]

        filename = file.filename.lower()
        file_extension = None

        for ext in allowed_extensions:
            if filename.endswith(ext):
                file_extension = ext
                break

        if not file_extension:
            raise HTTPException(status_code=400, detail="不支持的文件格式，支持: txt, md, docx, pdf")

        content = await file.read()
        print(f"[INFO] File size: {len(content)} bytes")
        text_content = ""

        if file_extension == ".txt" or file_extension == ".md":
            text_content = content.decode("utf-8", errors="ignore")
            print(f"[INFO] Extracted text length: {len(text_content)} characters")

        elif file_extension == ".docx":
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(content))
                
                # 提取段落文本
                paragraphs = [para.text for para in doc.paragraphs]
                text_content = "\n".join([p for p in paragraphs if p.strip()])
                print(f"[INFO] Extracted {len(paragraphs)} paragraphs from DOCX")
                print(f"[INFO] Extracted text length from DOCX: {len(text_content)} characters")
                
                # 如果段落提取为空，尝试从表格中提取
                if not text_content.strip():
                    print("[INFO] No text from paragraphs, trying tables...")
                    table_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                table_text.append(cell.text)
                    text_content = "\n".join([t for t in table_text if t.strip()])
                    print(f"[INFO] Extracted {len(table_text)} cells from tables")
                    print(f"[INFO] Text length from tables: {len(text_content)} characters")
                
                # 尝试使用 docx2txt 作为备选
                if not text_content.strip():
                    print("[INFO] Trying docx2txt as fallback...")
                    try:
                        import docx2txt
                        text_content = docx2txt.process(io.BytesIO(content))
                        print(f"[INFO] Extracted text length from docx2txt: {len(text_content)} characters")
                    except ImportError:
                        print("[WARN] docx2txt not installed, skipping")
                    except Exception as e:
                        print(f"[WARN] docx2txt failed: {str(e)}")
                
            except ImportError:
                raise HTTPException(status_code=500, detail="需要安装 python-docx 库")
            except Exception as e:
                print(f"[ERROR] DOCX parsing error: {str(e)}")
                raise HTTPException(status_code=500, detail=f"解析 Word 文档失败: {str(e)}")

        elif file_extension == ".pdf":
            try:
                from PyPDF2 import PdfReader
                import io
                reader = PdfReader(io.BytesIO(content))
                text_content = "\n".join([page.extract_text() for page in reader.pages])
                print(f"[INFO] Extracted text length from PDF: {len(text_content)} characters")
            except ImportError:
                raise HTTPException(status_code=500, detail="需要安装 PyPDF2 库")
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"解析 PDF 文档失败: {str(e)}")

        if not text_content.strip():
            print("[WARN] No text extracted from document, saving file only")
            
            # 保存文件到 document 目录
            base_dir = os.path.dirname(os.path.abspath(__file__))
            document_dir = os.path.join(base_dir, "..", "document")
            document_dir = os.path.normpath(document_dir)
            os.makedirs(document_dir, exist_ok=True)
            
            file_path = os.path.join(document_dir, file.filename)
            with open(file_path, "wb") as f:
                f.write(content)
            
            print(f"[INFO] File saved to: {file_path}")
            
            return {
                "status": "success",
                "filename": file.filename,
                "file_size": len(content),
                "strategy": strategy,
                "module_count": 0,
                "modules": [],
                "warning": "无法从文件中提取文本内容，文件已保存但无法用于智能问答"
            }

        print(f"[INFO] Processing document with strategy: {strategy}, target_chunks: {target_chunks}")
        result = process_document(
            text=text_content,
            strategy=strategy,
            target_chunks=target_chunks
        )

        print(f"[INFO] Document processed successfully, {len(result)} modules generated")

        # 保存文件到 document 目录
        base_dir = os.path.dirname(os.path.abspath(__file__))
        document_dir = os.path.join(base_dir, "..", "document")
        document_dir = os.path.normpath(document_dir)
        os.makedirs(document_dir, exist_ok=True)
        
        file_path = os.path.join(document_dir, file.filename)
        with open(file_path, "wb") as f:
            f.write(content)
        
        print(f"[INFO] File saved to: {file_path}")
        
        # 重新加载 RAG 索引
        print(f"[INFO] Reloading RAG index...")
        rag_service.initialize()
        print(f"[INFO] RAG index reloaded")

        return {
            "status": "success",
            "filename": file.filename,
            "file_size": len(content),
            "strategy": strategy,
            "module_count": len(result),
            "modules": result
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] File upload failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文件上传处理失败: {str(e)}")


if __name__ == "__main__":
    import uvicorn
    print("[INFO] Initializing RAG service...")
    rag_service.initialize()
    uvicorn.run(app, host="127.0.0.1", port=8000)